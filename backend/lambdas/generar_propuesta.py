"""POST /iniciativas/{id}/propuesta — Genera propuesta de arquitectura Word.

Flujo:
  1. Lee servicios AWS identificados desde S3
  2. Usa AWS MCP (Model Context Protocol) para estimar costos con IA
  3. Llama a Claude para generar el contenido de cada sección
  4. Usa la plantilla Word de S3 para generar el documento final
  5. Guarda el .docx en S3 y retorna URL de descarga
"""
import json
import boto3
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from common.helpers import build_response, DynamoHelper, S3Helper, log_event
from common.aws_pricing_mcp import estimar_costos_mcp

PLANTILLA_KEY = "plantillas/propuesta_arquitectura_template.docx"
REGION = "US East (N. Virginia)"


def handler(event, context):
    path_params = event.get("pathParameters") or {}
    id_iniciativa = path_params.get("id")
    if not id_iniciativa:
        return build_response(400, {"error": "id requerido"})

    meta = DynamoHelper.get_meta(id_iniciativa)
    if not meta:
        return build_response(404, {"error": "Iniciativa no encontrada"})

    if meta.get("estado") not in ("DISCOVERY_OK", "EXTRAIDO", "EN_REVISION", "VALIDADO", "ERROR_PROPUESTA"):
        return build_response(400, {"error": "La iniciativa aún no tiene validación completada. Debes validar los documentos en el Paso 1 antes de generar la propuesta.", "estado_actual": meta.get("estado")})

    DynamoHelper.update_meta(id_iniciativa, {"estado": "GENERANDO_PROPUESTA"})

    # Invocar worker asíncrono
    lambda_client = boto3.client("lambda", region_name="us-east-1")
    lambda_client.invoke(
        FunctionName="coppel-cloud-prod-propuesta-worker",
        InvocationType="Event",
        Payload=json.dumps({"id_iniciativa": id_iniciativa})
    )

    return build_response(202, {
        "message": f"Generación de propuesta iniciada para {id_iniciativa}",
        "estado": "GENERANDO_PROPUESTA"
    })


def _set_progreso(id_iniciativa: str, paso: int, total: int, mensaje: str):
    """Actualiza progreso en S3 para que el frontend lo lea."""
    S3Helper.put_json(f"iniciativas/{id_iniciativa}/progreso.json", {
        "paso": paso, "total": total, "mensaje": mensaje,
        "porcentaje": int((paso / total) * 100)
    })


def generar_propuesta_worker(id_iniciativa: str, meta: dict):
    """Worker principal — llamado desde propuesta_worker.py"""
    TOTAL = 6

    # 1. Leer servicios identificados
    _set_progreso(id_iniciativa, 1, TOTAL, "Leyendo servicios identificados...")
    try:
        servicios_raw = S3Helper.read_text(f"iniciativas/{id_iniciativa}/resultados/servicios_aws.json")
        servicios = json.loads(servicios_raw)
    except Exception:
        servicios = {"servicios": [], "resumen_arquitectura": ""}

    # 2. Leer discovery validado
    try:
        discovery_raw = S3Helper.read_text(f"iniciativas/{id_iniciativa}/resultados/validacion_discovery.json")
    except Exception:
        discovery_raw = "{}"

    # 3. Leer transcripciones
    insumos = meta.get("insumos") or []
    trans_texto = ""
    for i in insumos:
        if i.get("tipo") == "transcripciones":
            try:
                from urllib.parse import unquote_plus
                trans_texto += S3Helper.read_text(unquote_plus(i["key"]))
            except Exception:
                pass

    # 4. Consultar precios AWS via MCP (Claude como agente de pricing)
    _set_progreso(id_iniciativa, 2, TOTAL, "Consultando precios AWS con IA...")
    contexto_pricing = f"Proyecto: {meta.get('nombre', '')}. Ambiente: {meta.get('ambiente', 'Prod')}. "
    contexto_pricing += servicios.get("resumen_arquitectura", "")
    costos = estimar_costos_mcp(servicios.get("servicios", []), contexto_pricing)

    # 5. Llamar a Claude para generar contenido
    _set_progreso(id_iniciativa, 3, TOTAL, "Generando contenido con Claude (8 secciones)...")
    prompt_template = S3Helper.read_text("prompts/generar_propuesta.md")
    prompt = prompt_template \
        .replace("{DISCOVERY}", discovery_raw) \
        .replace("{TRANSCRIPCIONES}", trans_texto or "No disponibles") \
        .replace("{SERVICIOS}", json.dumps(servicios, ensure_ascii=False)) \
        .replace("{COSTOS}", json.dumps(costos, ensure_ascii=False))

    from common.claude_client import ClaudeClient
    respuesta = ClaudeClient().generate(
        system="Eres un arquitecto cloud senior de Coppel. Responde ÚNICAMENTE con JSON válido.",
        prompt=prompt
    )["text"]

    contenido = _parse_json(respuesta)

    # 6. Generar Word usando plantilla
    _set_progreso(id_iniciativa, 4, TOTAL, "Armando documento Word con plantilla...")
    docx_bytes = _generar_word(contenido, costos, meta)

    # 7. Guardar en S3
    _set_progreso(id_iniciativa, 5, TOTAL, "Guardando documento en S3...")
    output_key = f"iniciativas/{id_iniciativa}/salidas/propuesta_arquitectura.docx"
    S3Helper.put_bytes(output_key,
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # 7b. Guardar costos como JSON para el frontend
    S3Helper.put_json(f"iniciativas/{id_iniciativa}/resultados/costos_estimados.json", costos)

    # 8. Generar URL de descarga
    url_descarga = S3Helper.presigned_get(output_key, expires=86400)

    # 9. Actualizar estado
    _set_progreso(id_iniciativa, 6, TOTAL, "Finalizando...")
    salidas = dict(meta.get("salidas") or {})
    salidas["propuesta"] = output_key
    DynamoHelper.update_meta(id_iniciativa, {
        "estado": "GENERADO",
        "salidas": salidas,
    })
    DynamoHelper.put_event(id_iniciativa, "PROPUESTA_GENERADA", {"key": output_key})
    log_event("generar_propuesta", {"id_iniciativa": id_iniciativa})

    return url_descarga





def _generar_word(contenido: dict, costos: dict, meta: dict) -> bytes:
    """Genera el documento Word usando la plantilla y el contenido generado por Claude."""
    # Cargar plantilla
    plantilla_bytes = S3Helper.read_bytes(PLANTILLA_KEY)
    doc = Document(BytesIO(plantilla_bytes))

    proyecto = contenido.get("proyecto", {})
    secciones = contenido.get("secciones", {})

    # Limpiar contenido existente manteniendo estilos
    _limpiar_documento(doc)

    # Portada
    _agregar_portada(doc, proyecto, meta)

    # Tabla de contenido
    doc.add_heading("Contenido", level=1)
    secciones_toc = [
        "1. Introducción", "2. Requerimiento Técnico", "3. Premisas",
        "4. Arquitectura de infraestructura AWS", "5. Fuera de alcance",
        "6. Consideraciones generales", "7. Propuesta económica AWS",
        "8. Consideraciones comerciales"
    ]
    for s in secciones_toc:
        doc.add_paragraph(s)

    # 1. Introducción
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(secciones.get("introduccion", ""))

    # 2. Requerimiento Técnico
    doc.add_heading("2. Requerimiento Técnico", level=1)
    doc.add_paragraph(secciones.get("requerimiento_tecnico", ""))

    # 3. Premisas
    doc.add_heading("3. Premisas", level=1)
    for premisa in (secciones.get("premisas") or []):
        _add_bullet(doc, premisa)

    # 4. Arquitectura AWS
    doc.add_heading("4. Arquitectura de infraestructura AWS", level=1)
    doc.add_paragraph(secciones.get("arquitectura_aws", ""))

    # 5. Fuera de alcance
    doc.add_heading("5. Fuera de alcance", level=1)
    for item in (secciones.get("fuera_de_alcance") or []):
        _add_bullet(doc, item)

    # 6. Consideraciones generales
    doc.add_heading("6. Consideraciones generales", level=1)
    for item in (secciones.get("consideraciones_generales") or []):
        _add_bullet(doc, item)

    # 7. Propuesta económica
    doc.add_heading("7. Propuesta económica AWS", level=1)
    _agregar_tabla_costos(doc, costos)

    # 8. Consideraciones comerciales
    doc.add_heading("8. Consideraciones comerciales", level=1)
    for item in (secciones.get("consideraciones_comerciales") or []):
        _add_bullet(doc, item)

    # Sección de aprobación
    doc.add_heading("Aprobación", level=1)
    doc.add_paragraph("El presente documento es aprobado por las siguientes personas:")

    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_bullet(doc: Document, text: str):
    """Agrega un párrafo con bullet."""
    styles = [s.name for s in doc.styles if s.name and ('bullet' in s.name.lower() or 'list' in s.name.lower())]
    if styles:
        p = doc.add_paragraph(style=styles[0])
        p.add_run(text)
    else:
        doc.add_paragraph(f"• {text}")


def _limpiar_documento(doc: Document):
    """Elimina párrafos del body manteniendo headers/footers/estilos."""
    for para in doc.paragraphs[5:]:  # Mantener primeros párrafos de portada
        para.clear()


def _agregar_portada(doc: Document, proyecto: dict, meta: dict):
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.clear()
    p.add_run("Proyecto:").bold = True
    doc.add_paragraph(proyecto.get("nombre", meta.get("nombre", "")))
    doc.add_paragraph(proyecto.get("ambiente", meta.get("ambiente", "")))
    doc.add_paragraph(proyecto.get("unidad_negocio", meta.get("solicitante", "")))


def _agregar_tabla_costos(doc: Document, costos: dict):
    servicios = costos.get("servicios_cotizados", [])
    if not servicios:
        doc.add_paragraph("Estimación de costos pendiente.")
        return

    tabla = doc.add_table(rows=1, cols=5)
    try:
        tabla.style = "Table Grid"
    except KeyError:
        pass

    # Header
    hdr = tabla.rows[0].cells
    hdr[0].text = "Servicio"
    hdr[1].text = "Especificación"
    hdr[2].text = "Categoría"
    hdr[3].text = "Prioridad"
    hdr[4].text = "Costo mensual (USD)"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Filas
    for svc in servicios:
        row = tabla.add_row().cells
        row[0].text = svc.get("servicio", "")
        row[1].text = svc.get("especificacion", svc.get("nota", ""))
        row[2].text = svc.get("categoria", "")
        row[3].text = svc.get("prioridad", "")
        precio = svc.get("precio_mensual_usd")
        row[4].text = f"${precio:,.2f}" if precio else "A consultar"

    # Total
    total_row = tabla.add_row().cells
    total_row[0].text = "TOTAL MENSUAL"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[4].text = f"${costos.get('total_mensual_usd', 0):,.2f} USD"
    total_row[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph(f"\nEquivalente anual (12 meses): ${costos.get('total_anual_usd', 0):,.2f} USD")

    # Supuestos de la estimación
    supuestos = costos.get("supuestos", [])
    if supuestos:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Supuestos de la estimación:").bold = True
        for s in supuestos:
            doc.add_paragraph(f"• {s}")

    # Recomendaciones de ahorro
    recomendaciones = costos.get("recomendaciones_ahorro", [])
    if recomendaciones:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Recomendaciones de optimización de costos:").bold = True
        for r in recomendaciones:
            doc.add_paragraph(f"• {r}")


def _parse_json(texto: str) -> dict:
    try:
        return json.loads(texto)
    except Exception:
        match = re.search(r"\{.*\}", texto, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except Exception:
                pass
    return {}
